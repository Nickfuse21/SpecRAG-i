"""
Step 3 of ingestion: turn a .docx into a clause tree.

This is the hardest part of the pipeline and the one that decides whether
retrieval works at all. A 3GPP spec is not prose — it is a numbered clause
tree with prose, tables, and ASN.1 code all interleaved. Get the structure
wrong here and every downstream stage inherits the damage.

What we learned inspecting a real file (38331, RRC) before writing this:

  - Body clauses use paragraph styles "Heading 1".."Heading 7", and the
    heading text is "<number>\t<title>", e.g. "5.3.5.3\tReception of an
    RRCReconfiguration by the UE".
  - Annex banners ("Annex A (informative):\tGuidelines...") use a SEPARATE
    style, "Heading 8" — not part of the 1-7 numbering. They start a new,
    independent top-level tree (annex sub-clauses restart at depth 1 as
    "A.1", "A.2.1", ...).
  - ASN.1 has its own dedicated paragraph style, "PL" (Program Listing).
    We do NOT need a heuristic (monospace font, "::=" regex) — the style
    name is a direct, reliable signal.
  - "toc 1".."toc 5" styles are the auto-generated table of contents and
    must be skipped entirely — they duplicate heading text and would
    otherwise pollute the tree.
  - "Void" clauses are deleted/reserved clause numbers with no content.
    They are dropped.

Critically: we walk `doc.element.body` in DOCUMENT ORDER, not
`doc.paragraphs` and `doc.tables` separately. Those are two disconnected
lists in python-docx — iterating them separately loses interleaving, and a
table silently detaches from the clause it belongs to. This is the most
common bug in a 3GPP parser and it is silent: nothing crashes, the output
just quietly attributes tables to the wrong (or no) clause.

Usage
-----
    python -m src.ingest.parse_docx                  # all .docx in data/docx/
    python -m src.ingest.parse_docx --file 38331-ia0.docx
    python -m src.ingest.parse_docx --debug-clause 5.3.5.3 --file 38331-ia0.docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import tiktoken
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))
from src.config import DOCX_DIR, DROP_CLAUSE_TITLES, PARSED_DIR, TOKENIZER_ENCODING  # noqa: E402
from src.ingest.download import decode_version  # noqa: E402

_ENC = tiktoken.get_encoding(TOKENIZER_ENCODING)


def count_tokens(text: str) -> int:
    return len(_ENC.encode(text, disallowed_special=()))


# --------------------------------------------------------------------------
# heading parsing
# --------------------------------------------------------------------------
_NUMBERED_HEADING = re.compile(
    r"^([A-Za-z0-9]+(?:\.[A-Za-z0-9]+)*)\s+(.+)$", re.DOTALL
)
_ANNEX_BANNER = re.compile(
    r"^Annex\s+([A-Za-z0-9]+)\s*\(([^)]*)\)\s*:\s*(.+)$", re.IGNORECASE | re.DOTALL
)


def normalise_heading_text(raw: str) -> str:
    """Docx headings separate number/title with a tab, occasionally a newline."""
    return re.sub(r"\s+", " ", raw.replace("\t", " ").replace("\n", " ")).strip()


@dataclass
class ParsedHeading:
    clause_id: str | None   # None for un-numbered headings (Foreword, ...)
    title: str
    level: int
    is_annex_banner: bool = False
    annex_kind: str = ""    # "informative" / "normative", banners only


def parse_heading(style_name: str, raw_text: str) -> ParsedHeading:
    text = normalise_heading_text(raw_text)

    # Trust the docx style number as ground-truth depth ("Heading 4" -> 4).
    # We used to derive depth by counting dots in the clause number
    # (e.g. "A.2.1".count(".") + 1 = 3) but that is wrong for annex
    # sub-clauses: the leading letter segment ("A") doesn't add a real
    # depth increment, so "A.2.1" is actually styled Heading 2, not
    # Heading 3. Real-data check against 38331-ia0.docx confirmed this:
    # "A.2.1" -> Heading 2, "A.3.1.1" -> Heading 3.
    try:
        style_level = int(style_name.rsplit(" ", 1)[-1])
    except ValueError:
        style_level = 1

    if style_level == 8:
        m = _ANNEX_BANNER.match(text)
        if m:
            letter, kind, title = m.groups()
            return ParsedHeading(
                clause_id=f"Annex {letter}", title=title.strip(), level=1,
                is_annex_banner=True, annex_kind=kind.strip().lower(),
            )
        # fallback: treat the whole banner as an unparsed top-level section
        return ParsedHeading(clause_id=text, title=text, level=1, is_annex_banner=True)

    m = _NUMBERED_HEADING.match(text)
    if m:
        clause_id, title = m.groups()
        return ParsedHeading(clause_id=clause_id, title=title.strip(), level=style_level)

    # un-numbered heading: "Foreword" (Heading 1, genuinely top-level), but
    # also Annex B's 727 "– <IEName>" entries, which are Heading 4 despite
    # having no clause number. Use the style's own depth, not a hardcoded 1,
    # and strip the leading bullet/dash so the clause_id is clean.
    clean = re.sub(r"^[–—\-•]\s*", "", text).strip()
    clean = clean or text
    return ParsedHeading(clause_id=clean, title=clean, level=style_level)


# --------------------------------------------------------------------------
# clause tree
# --------------------------------------------------------------------------
@dataclass
class Clause:
    clause_id: str
    title: str
    level: int
    breadcrumb: str
    blocks: list[dict] = field(default_factory=list)   # [{"type": "...", "text": "..."}]

    def add_block(self, block_type: str, text: str) -> None:
        if not text.strip():
            return
        # merge consecutive same-type blocks so prose paragraphs don't
        # fragment into dozens of tiny list entries
        if self.blocks and self.blocks[-1]["type"] == block_type:
            self.blocks[-1]["text"] += "\n" + text
        else:
            self.blocks.append({"type": block_type, "text": text})

    @property
    def full_text(self) -> str:
        return "\n\n".join(b["text"] for b in self.blocks)

    @property
    def token_count(self) -> int:
        return count_tokens(self.full_text)


# --------------------------------------------------------------------------
# table -> markdown
# --------------------------------------------------------------------------
def table_to_markdown(table: Table) -> str:
    """
    GFM pipe table. Header row kept (it's what makes a table self-describing
    once it lands in an LLM's context, detached from its page).

    Handles horizontal merges (python-docx repeats the same Cell object
    across a merged span; we de-dupe by identity of the underlying XML
    element). Vertical merges are NOT de-duplicated — a small, accepted
    limitation for a 3-day build; the repeated text is redundant but not
    harmful to retrieval.
    """
    rows: list[list[str]] = []
    for row in table.rows:
        cells: list[str] = []
        seen_tc = None
        for cell in row.cells:
            if cell._tc is seen_tc:
                continue  # horizontal-merge continuation, already emitted
            seen_tc = cell._tc
            text = cell.text.replace("\r", " ").replace("\n", "<br>").strip()
            text = text.replace("|", "\\|")
            cells.append(text)
        if any(c for c in cells):
            rows.append(cells)

    if not rows:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


# --------------------------------------------------------------------------
# main walk
# --------------------------------------------------------------------------
def iter_block_items(doc: Document):
    """
    Yield Paragraph / Table objects from doc.element.body IN DOCUMENT ORDER.

    This is the fix for the #1 bug in DIY 3GPP parsers: doc.paragraphs and
    doc.tables are two separate lists in python-docx. Iterating them
    separately loses interleaving, so a table silently detaches from the
    clause it appeared under.
    """
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield Paragraph(child, doc)
        elif tag == "tbl":
            yield Table(child, doc)
        # sectPr and anything else: not content, skip


DROP_TITLE_SET = {t.lower() for t in DROP_CLAUSE_TITLES}


def should_drop(title: str) -> bool:
    t = title.strip().lower()
    return t in DROP_TITLE_SET or t == "void"


# Some clauses keep their heading but carry a placeholder body instead of the
# title "Void" — 21.905 does this for unused alphabet index letters ("<void>").
# They are real clauses with zero information; indexing them just gives the
# retriever near-empty vectors that match nothing useful.
_VOID_BODY = {"void", "<void>", "(void)", "-", "–"}


def is_void_body(text: str) -> bool:
    return text.strip().lower() in _VOID_BODY


def parse_document(path: Path, spec: str, version: str, release: int) -> list[Clause]:
    doc = Document(str(path))

    clauses: list[Clause] = []
    stack: list[Clause] = []          # ancestor chain, current clause is stack[-1]
    current: Clause | None = None

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            style = item.style.name if item.style else "Normal"

            if style.startswith("toc"):
                continue  # auto-generated table of contents — noise, not content

            if style.startswith("Heading"):
                h = parse_heading(style, item.text)
                if h.clause_id is None and not item.text.strip():
                    continue

                if h.is_annex_banner:
                    stack = []  # annexes are a fresh top-level tree

                while stack and stack[-1].level >= h.level:
                    stack.pop()

                breadcrumb = " > ".join(f"{c.clause_id or c.title} {c.title}" for c in stack)
                clause_id = h.clause_id or h.title
                current = Clause(
                    clause_id=clause_id, title=h.title, level=h.level, breadcrumb=breadcrumb
                )
                clauses.append(current)
                stack.append(current)
                continue

            if current is None:
                continue  # front-matter before the first heading (title page etc.)

            text = item.text.strip()
            if not text:
                continue

            if style == "PL":
                current.add_block("asn1", text)
            else:
                current.add_block("prose", text)

        elif isinstance(item, Table):
            if current is None:
                continue
            md = table_to_markdown(item)
            if md:
                current.add_block("table", md)

    return [
        c for c in clauses
        if not should_drop(c.title)
        and c.full_text.strip()
        and not is_void_body(c.full_text)
    ]


# --------------------------------------------------------------------------
def stem_to_spec(stem: str) -> tuple[str, str, str, int]:
    """'38331-ia0' -> ('38.331', 'ia0', '18.10.0', 18)"""
    num, code = stem.split("-")
    spec = f"{num[:2]}.{num[2:]}"
    x, y, z = decode_version(code)
    return spec, code, f"{x}.{y}.{z}", x


def write_jsonl(clauses: list[Clause], meta: dict, out_path: Path) -> None:
    # A clause_id is NOT unique within a spec. 21.905 (Vocabulary) has an
    # alphabet index heading "A", "B", "C"... once under Definitions and again
    # under Abbreviations. The clause_id stays exactly as the document says it
    # (citations must be truthful), but the UID we index by gets an occurrence
    # suffix so it can serve as a primary key.
    seen: dict[str, int] = {}

    with open(out_path, "w", encoding="utf-8") as f:
        for i, c in enumerate(clauses):
            n = seen.get(c.clause_id, 0) + 1
            seen[c.clause_id] = n
            uid = f"{meta['stem']}::{c.clause_id}" + (f"#{n}" if n > 1 else "")
            record = {
                "clause_uid": uid,
                "spec_id": f"TS {meta['spec']}",
                "spec": meta["spec"],
                "version": meta["version"],
                "release": meta["release"],
                "clause_id": c.clause_id,
                "clause_title": c.title,
                "level": c.level,
                "breadcrumb": c.breadcrumb,
                "blocks": c.blocks,
                "token_count": c.token_count,
                "order": i,
            }
            f.write(json.dumps(record, ensure_ascii=False) + "\n")


def main() -> int:
    ap = argparse.ArgumentParser(description="Parse .docx specs into clause-tree JSONL")
    ap.add_argument("--file", action="append", default=None,
                    help="parse only this filename (repeatable), e.g. --file 38331-ia0.docx")
    ap.add_argument("--debug-clause", default=None,
                    help="print the full parsed record for one clause_id and exit")
    args = ap.parse_args()

    files = [DOCX_DIR / f for f in args.file] if args.file else sorted(DOCX_DIR.glob("*.docx"))
    if not files:
        print(f"No .docx files in {DOCX_DIR}")
        return 1

    total_clauses = 0
    total_tokens = 0

    for path in files:
        stem = path.stem
        try:
            spec, code, version, release = stem_to_spec(stem)
        except Exception as e:                                          # noqa: BLE001
            print(f"! skipping {path.name}: can't parse stem ({e})")
            continue

        print(f"{spec}  ({path.name})")
        clauses = parse_document(path, spec, version, release)

        if args.debug_clause:
            hit = next((c for c in clauses if c.clause_id == args.debug_clause), None)
            if hit:
                print(json.dumps({
                    "clause_id": hit.clause_id, "title": hit.title,
                    "breadcrumb": hit.breadcrumb, "blocks": hit.blocks,
                    "token_count": hit.token_count,
                }, indent=2, ensure_ascii=False))
                return 0
            print(f"  clause {args.debug_clause!r} not found in {stem}")
            continue

        meta = {"stem": stem, "spec": spec, "version": version, "release": release}
        out_path = PARSED_DIR / f"{stem}.jsonl"
        write_jsonl(clauses, meta, out_path)

        n_tokens = sum(c.token_count for c in clauses)
        n_asn1 = sum(1 for c in clauses for b in c.blocks if b["type"] == "asn1")
        n_tables = sum(1 for c in clauses for b in c.blocks if b["type"] == "table")
        print(f"  -> {len(clauses)} clauses, {n_tokens:,} tokens, "
              f"{n_asn1} asn1 blocks, {n_tables} table blocks")
        print(f"  -> {out_path.name}")

        total_clauses += len(clauses)
        total_tokens += n_tokens

    if not args.debug_clause:
        print(f"\nDone. {total_clauses} clauses, {total_tokens:,} tokens across {len(files)} spec(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
